"""Tools del Docs Agent: documentos Word desde contenido de Looker.

Dos modos: exportar un dashboard como reporte (título + sección por tile con
imagen y muestra de datos) o componer un documento narrativo por secciones.
El template corporativo (.docx en el bucket) aporta estilos y portada.
"""
import io
import json

from docx import Document
from docx.shared import Inches, Pt

from common.delivery import SIGNED_URL_HOURS, timestamped_name, upload_and_sign
from common.looker_client import get_sdk, poll_render_task

IMG_W, IMG_H = 1200, 700
DATA_SAMPLE_ROWS = 10


def _base_document(template_name: str = "") -> Document:
    if not template_name:
        return Document()
    from common import templates
    name = template_name if template_name.endswith(".docx") else f"{template_name}.docx"
    return Document(io.BytesIO(templates.load_bytes("docs", name)))


def _add_data_table(doc: Document, rows: list[dict], max_rows: int = DATA_SAMPLE_ROWS):
    if not rows:
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h.split(".")[-1].replace("_", " ").title()
        cell.paragraphs[0].runs[0].bold = True
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for c, h in enumerate(headers):
            v = row.get(h)
            cells[c].text = "" if v is None else (f"{v:,.2f}" if isinstance(v, float) else str(v))
    if len(rows) > max_rows:
        p = doc.add_paragraph(f"… {len(rows) - max_rows} filas adicionales en el dataset completo.")
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)


def export_dashboard_to_docx(dashboard_id: str, filename: str = "",
                             template_name: str = "", include_data_sample: bool = True) -> str:
    """Genera un reporte .docx desde un dashboard: título + una sección por
    tile con su visualización como imagen y (opcional) una muestra de datos."""
    sdk = get_sdk()
    db = sdk.dashboard(dashboard_id=dashboard_id, fields="id,title,dashboard_elements")
    doc = _base_document(template_name)
    doc.add_heading(db.title or "Reporte", level=0)
    exported, skipped = [], []
    for el in (db.dashboard_elements or []):
        if el.type != "vis" or not el.query_id:
            continue
        doc.add_heading(el.title or f"Tile {el.id}", level=1)
        try:
            task = sdk.create_query_render_task(query_id=el.query_id,
                                                result_format="png",
                                                width=IMG_W, height=IMG_H)
            doc.add_picture(io.BytesIO(poll_render_task(sdk, task.id)), width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph(f"(visualización no disponible: {e})")
            skipped.append(el.title or el.id)
        if include_data_sample:
            try:
                raw = sdk.run_query(query_id=el.query_id, result_format="json")
                _add_data_table(doc, json.loads(raw) if isinstance(raw, str) else raw)
            except Exception:
                pass
        exported.append(el.title or el.id)
    if not exported:
        return json.dumps({"error": "El dashboard no tiene tiles exportables."})
    buf = io.BytesIO()
    doc.save(buf)
    url = upload_and_sign(buf.getvalue(),
                          timestamped_name(filename or db.title or "reporte", ".docx"))
    return json.dumps({"download_url": url, "sections": exported, "skipped": skipped,
                       "expires_in_hours": SIGNED_URL_HOURS}, ensure_ascii=False)


def create_document(title: str, sections: list[dict], filename: str = "",
                    template_name: str = "") -> str:
    """Compone un documento narrativo. Cada sección: {"heading": str, "body": str}.
    Útil para resúmenes ejecutivos redactados por el orquestador con cifras
    validadas por el Catalog."""
    doc = _base_document(template_name)
    doc.add_heading(title, level=0)
    for s in sections:
        doc.add_heading(s.get("heading", ""), level=1)
        for para in (s.get("body") or "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    url = upload_and_sign(buf.getvalue(), timestamped_name(filename or title, ".docx"))
    return json.dumps({"download_url": url, "sections": len(sections),
                       "expires_in_hours": SIGNED_URL_HOURS})


def list_docs_templates() -> str:
    """Lista los templates .docx corporativos disponibles."""
    from common import templates
    return json.dumps(templates.list_names("docs"))


ALL_TOOLS = [export_dashboard_to_docx, create_document, list_docs_templates]
